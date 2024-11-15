# data_processing.py

import os
import pandas as pd
import datetime
import shutil
import docx
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from bs4 import BeautifulSoup
import json
from typing import List
from jinja2 import Environment, FileSystemLoader
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def create_working_dir(working_dir: str, name: str) -> str:
    """
    Create a working directory with a timestamp and sanitized name.

    Parameters
    ----------
    working_dir: Base directory where the new directory will be created.
    name: Name to be sanitized and included in the directory name.

    Returns
    -------
    Path to the newly created directory.
    """
    date = (
        str(datetime.datetime.now())
        .split(".")[0]
        .replace(" ", "_")
        .replace("-", "_")
        .replace(":", "_")
    )
    name = (
        name.replace("-", "_")
        .replace(":", "_")
        .replace("(", "_")
        .replace(")", "_")
    )
    dir_name = os.path.join(working_dir, f"{date}_{name}")
    if not os.path.exists(dir_name):
        os.mkdir(dir_name)
    return dir_name


def sanitize_filename(filename: str) -> str:
    """
    Sanitize a filename by replacing or removing invalid characters.

    Parameters
    ----------
    filename: Filename to sanitize.

    Returns
    -------
    Sanitized filename.
    """
    return (
        filename.replace(" ", "_")
        .replace("-", "_")
        .replace("/", "_")
        .replace("\\", "_")
        .replace("*", "")
        .replace("|", "")
        .replace(":", "")
        .replace("?", "")
        .replace("<", "")
        .replace(">", "")
        .replace(":", "_")
        .replace("(", "_")
        .replace(")", "_")
    )


def word_cv_prepare(
    working_dir: str, save_dir: str, skills: List[str], position: str,
    current_skills: List[str]
) -> None:
    """
    Prepare and save a CV document in a specified directory.

    Parameters
    ----------
    working_dir: Base working directory.
    save_dir: Directory where the CV will be saved.
    skills: List of skills to include in the CV.
    position: Job position to include in the CV.
    current_skills: List of current skills to include in the CV.

    Returns
    -------
    None
    """
    all_skills = sorted(set(current_skills + skills), key=str.lower)

    source = os.path.join(working_dir, "PT.docx")
    if not os.path.exists(source):
        print(f"Source CV file not found: {source}")
        return

    destination = os.path.join(save_dir, "PrzemyslawTuturCV.docx")
    try:
        shutil.copy(source, destination)
    except OSError as e:
        print(f"Error copying document: {e}")
        return

    doc_name_position = sanitize_filename(position)
    try:
        doc = docx.Document(destination)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # Add all skills
    paragraph = doc.add_paragraph(", ".join(all_skills).upper())
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = paragraph.runs[0]
    run.font.name = "Times New Roman"

    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para_run = footer_para.add_run(
        f"This CV document was automatically generated and submitted for the "
        f"{position} position based on skill matching. If you contact me with "
        f"a response, I might be momentarily confused :)... Apologies for the "
        f"Monty Python 'spam, spam, spam' scenario if you have received "
        f"multiple CVs."
    )
    footer_para_run.font.name = "Times New Roman"
    footer_para_run.font.bold = True
    footer_para_run.font.size = Pt(8)

    # Read certification and GitHub files
    certs_file = os.path.join(working_dir, "certs.txt")
    github_file = os.path.join(working_dir, "github.txt")

    certs_links = []
    github_links = []

    if os.path.exists(certs_file):
        with open(certs_file, "r") as file:
            certs_links = [line.strip() for line in file]

    if os.path.exists(github_file):
        with open(github_file, "r") as file:
            github_links = [line.strip() for line in file]

    # Add certifications
    if certs_links:
        paragraph2 = doc.add_paragraph("CERTIFICATES - LAST 3 YEARS")
        paragraph2.style = "Heading 1"
        paragraph2.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("\n" + "\n".join(certs_links))

    # Add GitHub links
    if github_links:
        paragraph3 = doc.add_paragraph("LINKEDIN/GITHUB")
        paragraph3.style = "Heading 1"
        paragraph3.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph("\n" + "\n".join(github_links) + "\n")

    # Add soft skills
    soft_skills = [
        "ADAPTABILITY",
        "TIME MANAGEMENT",
        "TEAM LEADERSHIP",
        "COMMUNICATION",
        "PROBLEM SOLVING"
    ]
    paragraph4 = doc.add_paragraph("SOFT SKILLS")
    paragraph4.style = "Heading 1"
    paragraph4.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("\n" + ", ".join(soft_skills) + "\n")

    # Add hobbies
    hobbies = ["READING", "MOTORCYCLING", "CLIMBING", "MARTIAL ARTS"]
    paragraph5 = doc.add_paragraph("HOBBIES")
    paragraph5.style = "Heading 1"
    paragraph5.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph5a = doc.add_paragraph("\n" + ", ".join(hobbies))
    paragraph5a.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Save the document
    try:
        doc.save(os.path.join(
            save_dir, f"Przemyslaw_Tutur_{doc_name_position}.docx")
        )
    except OSError as e:
        print(f"Error saving document: {e}")

    # Append - PROJECTS.docx with custom formatting based on tags
    projects_path = os.path.join(working_dir, "PROJECTS.docx")
    if os.path.exists(projects_path):
        projects_doc = docx.Document(projects_path)
        doc.add_page_break()

        for para in projects_doc.paragraphs:
            text = para.text.strip()

            if text.startswith("<main-info>"):
                # Main section - "MAIN PROJECTS"
                main_info_paragraph = doc.add_paragraph(text.replace("<main-info>", "").strip())
                main_info_paragraph.style = "Heading 1"
            elif text.startswith("<company>"):
                # Company name- bolded
                company_paragraph = doc.add_paragraph()
                company_run = company_paragraph.add_run(text.replace("<company>", "").strip())
                company_run.bold = True
                company_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            elif text.startswith("<project>"):
                # Project - tabulation + bold
                project_paragraph = doc.add_paragraph(text.replace("<project>", "").strip())
                project_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                project_paragraph.paragraph_format.first_line_indent = docx.shared.Pt(0)
                project_run = project_paragraph.runs[0]
                project_run.bold = True
            elif text.startswith("<project-desc>"):
                # Project description - double tabulation
                project_desc_paragraph = doc.add_paragraph(text.replace("<project-desc>", "").strip())
                project_desc_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                project_desc_paragraph.paragraph_format.first_line_indent = docx.shared.Pt(0)
            elif text.startswith("<project-skills>"):
                # Skills - double tabulation
                project_skills_paragraph = doc.add_paragraph(text.replace("<project-skills>", "").strip())
                project_skills_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                project_skills_paragraph.paragraph_format.first_line_indent = docx.shared.Pt(36)
            else:
                # other possibilities - if exists
                other_paragraph = doc.add_paragraph(text)
                other_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    try:
        doc.save(os.path.join(save_dir, f"Przemyslaw_Tutur_{doc_name_position}_extended.docx"))
    except OSError as e:
        print(f"Error saving document: {e}")


def generate_cover_letter(
        working_dir: str, job_title: str, company_name: str, job_url: str,
        skills: List[str], soft_skills: List[str]
) -> None:
    """
    Generate and save a cover letter for a job application.

    Parameters
    ----------
    working_dir: Base working directory.
    job_title: Job title to include in the cover letter.
    company_name: Company name to include in the cover letter.
    job_url: Job URL to include in the cover letter.
    skills: List of hard skills to include in the cover letter.
    soft_skills: List of soft skills to include in the cover letter.

    Returns
    -------
    None
    """
    technical_skills = ", ".join(skills)
    soft_skills_str = ", ".join(soft_skills)

    doc = docx.Document()
    doc.add_heading("Cover Letter", 0)
    p1 = doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    p1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    doc.add_paragraph("Dear Hiring Manager,")

    # Initial paragraphs
    doc.add_paragraph(
        f"I am writing to express my interest in the {job_title} position at "
        f"{company_name}. I found this job listing on {job_url} and believe "
        f"that my skills and experience make me a strong candidate for this "
        f"role."
    )

    # Paragraph for technical skills
    p = doc.add_paragraph(
        f"I have extensive experience in the required technical skills mentioned in the "
        f"job description, including {technical_skills}. I am confident that my "
        f"background and knowledge will enable me to contribute effectively "
        f"to your team."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Paragraph for soft skills
    p = doc.add_paragraph(
        f"In addition to my technical expertise, I possess strong soft skills such as {soft_skills_str}. "
        f"These skills have enabled me to work collaboratively and effectively in team environments, "
        f"and to manage time and projects efficiently."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Additional paragraphs with justified alignment
    p = doc.add_paragraph(
        "I look forward to the opportunity to discuss how my skills and "
        "experiences align with the needs of your team. Thank you for "
        "considering my application."
    )
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Przemyslaw Tutur")

    # Add footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para_run = footer_para.add_run(
        f"This motivation letter was generated and submitted for the "
        f"{job_title} position. Please contact me directly if you wish to use it for any other position."
    )
    footer_para_run.font.name = "Times New Roman"
    footer_para_run.font.bold = True
    footer_para_run.font.size = Pt(8)

    cover_letter_path = os.path.join(
        working_dir, f"Cover_Letter_{sanitize_filename(job_title)}.docx"
    )
    doc.save(cover_letter_path)
    print(f"Cover letter saved to {cover_letter_path}")


def take_job_description(dir: str, url: str) -> None:
    """
    Retrieve and save the job description from a given URL.

    Parameters
    ----------
    dir: The directory where the job description will be saved.
    url: URL of the job description.

    Returns
    -------
    None
    """
    try:
        resp = requests.get(url)
        soup = BeautifulSoup(resp.text, "html.parser")
        target_div_content = soup.find("div", class_="MuiBox-root css-7nl6k4")
        if target_div_content:
            with open(
                os.path.join(dir, "job_description.txt"), "w", encoding="utf-8"
            ) as fdescriptor:
                fdescriptor.write(target_div_content.text)
        else:
            print("The specified div was not found.")
    except Exception as e:
        print(f"Error taking job description: {e}")


def request(
    working_dir: str, current_skills: List[str], url: str
) -> pd.DataFrame:
    """
    Send a request to a job listing URL, process the data, and save it.

    Parameters
    ----------
    working_dir: Base working directory.
    current_skills: List of current skills to match against job listings.
    url: URL of the job listings.

    Returns
    -------
    DataFrame containing the processed job data.
    """
    try:
        # Clear the current output data CSV file
        file_path = os.path.join(working_dir, "output_data.csv")
        if os.path.exists(file_path):
            os.remove(file_path)

        resp = requests.get(url)
        int_resp = requests.get(url)
        json_string = str(
            int_resp.text.split('{"pages":')[1].split('"meta":')[0].rstrip(",")
            + "}]"
        )
        data_set = json.loads(json_string)
        data_list = []
        for data in data_set[0]["data"]:
            sub_url = "https://justjoin.it/offers/" + data["slug"]
            row = {
                "TITLE": data["title"],
                "REQUIRED_SKILLS": str(data["requiredSkills"]),
                "ADDITIONAL_SKILLS": str(data["niceToHaveSkills"]),
                "WORKPLACE_TYPE": data["workplaceType"],
                "REMOTE_INTERVIEW": data["remoteInterview"],
                "URL": sub_url,
                "PAYMENT_FROM": str(data["employmentTypes"][0]["fromPln"]),
                "PAYMENT_TO": str(data["employmentTypes"][0]["toPln"]),
                "LOCATION": data.get("city", "Unknown"),
                "COMPANY": data.get("companyName", "Unknown"),
                "DATE": datetime.datetime.now().strftime("%Y-%m-%d"),
            }
            data_list.append(row)
            directory = create_working_dir(working_dir, data["slug"])
            take_job_description(directory, sub_url)
            word_cv_prepare(
                working_dir,
                directory,
                data["requiredSkills"],
                data["title"],
                current_skills,
            )
            soft_skills = [
                "adaptability",
                "time management",
                "team leadership",
                "communication",
                "problem solving"
            ]
            generate_cover_letter(
                directory,
                data["title"],
                data.get("companyName", "Unknown"),
                sub_url,
                data["requiredSkills"],
                soft_skills
            )
            print("Processed:", data["title"])
        df = pd.DataFrame(data_list)
        write_header = not os.path.exists(file_path)
        df.to_csv(file_path, mode="a", index=False, header=write_header)

        # Append to the output_whole.csv
        whole_file_path = os.path.join(working_dir, "output_whole.csv")
        df.to_csv(
            whole_file_path, mode="a", index=False,
            header=not os.path.exists(whole_file_path)
        )

        print("Data appended to CSV.")
        return df
    except Exception as e:
        print(f"Error processing request: {e}")
        return pd.DataFrame()  # Return empty DataFrame on error

def request(
    working_dir: str, current_skills: List[str], url: str, job_type: str
) -> pd.DataFrame:
    """
    Send a request to a job listing URL, process the data, and save it.

    Parameters
    ----------
    working_dir: Base working directory.
    current_skills: List of current skills to match against job listings.
    url: URL of the job listings.
    job_type: Job type to be added to the DataFrame.

    Returns
    -------
    DataFrame containing the processed job data.
    """
    try:
        # Clear the current output data CSV file
        file_path = os.path.join(working_dir, "output_data.csv")
        if os.path.exists(file_path):
            os.remove(file_path)

        resp = requests.get(url)
        int_resp = requests.get(url)
        json_string = str(
            int_resp.text.split('{"pages":')[1].split('"meta":')[0].rstrip(",")
            + "}]"
        )
        data_set = json.loads(json_string)
        data_list = []
        for data in data_set[0]["data"]:
            sub_url = "https://justjoin.it/offers/" + data["slug"]
            required_skills = str(data["requiredSkills"])
            match_percentage = skill_match_percentage(required_skills,
                                                      current_skills)
            row = {
                "TITLE": data["title"],
                "REQUIRED_SKILLS": str(data["requiredSkills"]),
                "ADDITIONAL_SKILLS": str(data["niceToHaveSkills"]),
                "WORKPLACE_TYPE": data["workplaceType"],
                "REMOTE_INTERVIEW": data["remoteInterview"],
                "URL": sub_url,
                "PAYMENT_FROM": str(data["employmentTypes"][0]["fromPln"]),
                "PAYMENT_TO": str(data["employmentTypes"][0]["toPln"]),
                "LOCATION": data.get("city", "Unknown"),
                "COMPANY": data.get("companyName", "Unknown"),
                "DATE": datetime.datetime.now().strftime("%Y-%m-%d"),
                "JOB_TYPE": job_type, # Add job type to the row
                "MATCH_PERCENTAGE": match_percentage
            }
            data_list.append(row)
            directory = create_working_dir(working_dir, data["slug"])
            take_job_description(directory, sub_url)
            word_cv_prepare(
                working_dir,
                directory,
                data["requiredSkills"],
                data["title"],
                current_skills,
            )
            soft_skills = [
                "adaptability",
                "time management",
                "team leadership",
                "communication",
                "problem solving"
            ]
            generate_cover_letter(
                directory,
                data["title"],
                data.get("companyName", "Unknown"),
                sub_url,
                data["requiredSkills"],
                soft_skills=soft_skills
            )
            print("Processed:", data["title"])
        df = pd.DataFrame(data_list)
        write_header = not os.path.exists(file_path)
        df.to_csv(file_path, mode="a", index=False, header=write_header)

        # Append to the output_whole.csv
        whole_file_path = os.path.join(working_dir, "output_whole.csv")
        df.to_csv(
            whole_file_path, mode="a", index=False,
            header=not os.path.exists(whole_file_path)
        )

        print("Data appended to CSV.")
        return df
    except Exception as e:
        print(f"Error processing request: {e}")
        return pd.DataFrame()  # Return empty DataFrame on error


def skill_match_percentage(
    required_skills: List[str], current_skills: List[str]
) -> float:
    """
    Calculate the percentage of matching skills between required skills
    and current skills.

    Parameters
    ----------
    required_skills: List of required skills for a job.
    current_skills: List of current skills of the user.

    Returns
    -------
    Percentage of matching skills.
    """
    if not required_skills:
        return 0.0

    required_skills_lower = [skill.lower() for skill in required_skills]
    current_skills_lower = [skill.lower() for skill in current_skills]

    matched_skills = set(required_skills_lower).intersection(
        set(current_skills_lower))
    return len(matched_skills) / len(required_skills_lower) * 100


def generate_summary(working_dir: str, df: pd.DataFrame,
                     skills_file: str) -> None:
    summary_path = os.path.join(working_dir, "summary.txt")
    with open(summary_path, "w", encoding="utf-8") as summary_file:
        for index, row in df.iterrows():
            job_title = row['TITLE']
            match_percentage = row['MATCH_PERCENTAGE']
            url = row['URL']
            required_skills = [skill.lower() for skill in
                               eval(row['REQUIRED_SKILLS'])] if isinstance(
                row['REQUIRED_SKILLS'], str) else []

            with open(skills_file, "r") as file:
                all_skills = [line.strip().lower() for line in file.readlines()]

            missing_skills = set(required_skills) - set(all_skills)
            missing_skills_str = ", ".join(missing_skills)

            summary_file.write(f"{job_title}\n")
            summary_file.write(
                f"Procent pasujacych skilli: {match_percentage}%\n")
            summary_file.write(
                "Lista skilli ktorych nie ma na liscie skills.txt: "
                f"{missing_skills_str}\n"
            )
            summary_file.write(f"URL: {url}\n")
            summary_file.write(f"----------------\n")


def take_job_description(dir: str, url: str) -> None:
    """
    Retrieve and save the job description from a given URL.

    Parameters
    ----------
    dir: Directory where the job description will be saved.
    url: URL of the job description.

    Returns
    -------
    None
    """
    try:
        resp = requests.get(url)
        soup = BeautifulSoup(resp.text, "html.parser")
        target_div_content = soup.find_all("script")

        json_data = None
        for script in target_div_content:
            if 'body' in script.text:
                json_data = script.text
                break

        if json_data:
            # Konwertuj JSON-encoded string do słownika
            data = json.loads(json_data)
            #data['props']['pageProps']['offer']['body']

            with open(
                os.path.join(dir, "job_description.txt"), "w", encoding="utf-8"
            ) as fdescriptor:
                fdescriptor.write(f"Job URL: {url}\n\n")
                fdescriptor.write(data['props']['pageProps']['offer']["title"])
                fdescriptor.write(
                    str(data['props']['pageProps']['offer']['companyName'])
                )
                fdescriptor.write(
                    str(data['props']['pageProps']['offer']["employmentTypes"])
                )
                fdescriptor.write(
                    str(data['props']['pageProps']['offer']['body'])
                )
                fdescriptor.write(
                    str(data['props']['pageProps']['offer']['experienceLevel'])
                )
        else:
            print("The specified div was not found.")
    except Exception as e:
        print(f"Error taking job description: {e}")
